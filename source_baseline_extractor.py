from __future__ import annotations

import re
from collections import OrderedDict
from pathlib import Path
from typing import Any

import zipfile

from docx import Document
from pypdf import PdfReader

def _check_qcv_property(path: str | Path) -> bool:
    """Check if DOCX has custom property qcv_generated=true."""
    try:
        with zipfile.ZipFile(str(path), "r") as z:
            if "docProps/custom.xml" not in z.namelist():
                return False
            from lxml import etree
            root = etree.fromstring(z.read("docProps/custom.xml"))
            for prop in root:
                if prop.get("name") == "qcv_generated":
                    val = prop[0].text if len(prop) else ""
                    return val and val.strip().lower() == "true"
    except Exception:
        pass
    return False


# Known section title variants used to map source headings to normalized QCV blocks.
SECTION_ALIASES = {
    "summary": {"summary", "profile", "about", "professional summary", "overview"},
    "skills": {"skills", "technical skills", "tech snapshot", "top skills", "core skills"},
    "experience": {"experience", "professional experience", "work experience"},
    "education": {"education", "academic background"},
    "certifications": {"certifications", "certificates", "licenses & certifications", "licenses"},
    "languages": {"languages"},
}

MONTH_RX = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
DATE_LINE_RE = re.compile(rf"(?:{MONTH_RX}\s+\d{{4}}|\d{{4}})\s*(?:-|–|—|to)?\s*(?:Present|Current|{MONTH_RX}\s+\d{{4}}|\d{{4}})?", re.I)
EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)
PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")
URL_RE = re.compile(r"https?://\S+|www\.\S+", re.I)
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/\S+", re.I)
HEADING_STYLE_RE = re.compile(r"heading\s*\d+", re.I)
BULLET_PREFIX_RE = re.compile(r"^[•\-–—▪◦●]\s*")
TECH_SPLIT_RE = re.compile(r"\s*[,;|/]\s*")

# Lightweight technology vocabulary for deterministic skill extraction and grouping.
KNOWN_TECH = {
    "c#", "java", "python", "typescript", "javascript", "go", "rust", "sql",
    ".net", "asp.net", "asp.net core", "react", "angular", "vue", "fastapi", "graphql",
    "postgresql", "mysql", "redis", "mongodb", "kafka", "rabbitmq",
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "github actions",
    "jenkins", "gitlab ci", "prometheus", "grafana", "rag", "llm", "microservices", "rest apis",
}


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove non-breaking spaces."""
    return re.sub(r"\s+", " ", (text or "").replace("\xa0", " ")).strip()


def _looks_like_heading(text: str) -> bool:
    """Check whether a line matches one of the known section heading aliases."""
    t = _clean_text(text).rstrip(":")
    if not t:
        return False
    low = t.lower()
    return any(low == alias for aliases in SECTION_ALIASES.values() for alias in aliases)


def _heading_key(text: str) -> str | None:
    """Return the normalized section key for a recognized heading."""
    low = _clean_text(text).rstrip(":").lower()
    for key, aliases in SECTION_ALIASES.items():
        if low in aliases:
            return key
    return None


def _iter_paragraphs_with_style(doc: Document):
    """Yield normalized paragraph and table-row blocks with basic style metadata."""
    for p in doc.paragraphs:
        text = _clean_text(p.text)
        if not text:
            continue
        style_name = getattr(getattr(p, "style", None), "name", "") or ""
        yield {"text": text, "style": style_name, "kind": "paragraph"}

    for tbl in doc.tables:
        for row in tbl.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = " ".join(_clean_text(p.text) for p in cell.paragraphs if _clean_text(p.text))
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                yield {"text": " | ".join(row_cells), "style": "Table", "kind": "table_row"}


def _extract_docx_links(doc: Document) -> list[str]:
    """Collect unique external hyperlinks from DOCX relationships."""
    links: list[str] = []
    seen: set[str] = set()
    rels = getattr(getattr(doc, "part", None), "rels", {})
    for rel in rels.values():
        target = getattr(rel, "target_ref", None)
        if isinstance(target, str) and (target.startswith("http://") or target.startswith("https://") or target.startswith("www.")):
            if target not in seen:
                seen.add(target)
                links.append(target)
    return links


def extract_from_docx(path: str | Path) -> dict[str, Any]:
    """Extract normalized text blocks, links, and coarse sections from a DOCX source."""
    doc = Document(str(path))
    blocks = list(_iter_paragraphs_with_style(doc))
    links = _extract_docx_links(doc)

    sections: dict[str, list[str]] = OrderedDict()
    preamble: list[str] = []
    current = "preamble"
    sections[current] = preamble

    # Split blocks into coarse sections using either explicit heading text or heading styles.
    for block in blocks:
        text = block["text"]
        style = block["style"]
        maybe_heading = _heading_key(text)
        if maybe_heading or HEADING_STYLE_RE.search(style):
            key = maybe_heading or _heading_key(text) or _clean_text(text).rstrip(":").lower().replace(" ", "_")
            current = key
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(text)

    full_text = "\n".join(b["text"] for b in blocks)
    return {
        "source_type": "docx",
        "source_path": str(path),
        "full_text": full_text,
        "blocks": blocks,
        "links": links,
        "sections": sections,
        "qcv_generated": _check_qcv_property(path),
    }


def extract_from_pdf(path: str | Path) -> dict[str, Any]:
    """Extract page text and detected URLs from a PDF source."""
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        pages.append({"page": i, "text": page.extract_text() or ""})
    full_text = "\n\n".join(p["text"] for p in pages)
    return {
        "source_type": "pdf",
        "source_path": str(path),
        "full_text": full_text,
        "pages": pages,
        "links": URL_RE.findall(full_text),
        "sections": {},
    }


def extract_baseline(path: str | Path) -> dict[str, Any]:
    """Dispatch baseline extraction based on source file type."""
    path = Path(path)
    if path.suffix.lower() == ".docx":
        return extract_from_docx(path)
    if path.suffix.lower() == ".pdf":
        return extract_from_pdf(path)
    raise ValueError(f"Unsupported source file: {path}")


def _find_contacts(text: str, links: list[str]) -> dict[str, Any]:
    """Extract basic contact details and keep at most LinkedIn plus one website."""
    email = EMAIL_RE.search(text)
    phone = None
    for m in PHONE_RE.finditer(text):
        cand = m.group(1)
        if len(re.sub(r"\D", "", cand)) >= 9:
            phone = cand.strip()
            break

    linkedin = ""
    website = ""
    for link in links + URL_RE.findall(text):
        if LINKEDIN_RE.search(link):
            linkedin = link
        elif not website:
            website = link

    out_links = []
    for item in [linkedin, website]:
        item = _clean_text(item)
        if item and item not in out_links:
            out_links.append(item)

    return {
        "email": email.group(0) if email else "",
        "phone": phone or "",
        "links": out_links,
    }


def _split_summary(lines: list[str]) -> list[str]:
    """Normalize summary lines into a compact bullet list."""
    items: list[str] = []
    for line in lines:
        line = BULLET_PREFIX_RE.sub("", _clean_text(line))
        if not line:
            continue
        if len(line) < 4:
            continue
        items.append(line)
    return items[:8]


def _tokenize_skills(lines: list[str]) -> list[str]:
    """Split skill text into distinct technology tokens using simple deterministic rules."""
    items: list[str] = []
    for line in lines:
        clean = BULLET_PREFIX_RE.sub("", _clean_text(line))
        if not clean:
            continue
        if ":" in clean and len(clean) < 80:
            _, right = clean.split(":", 1)
            clean = right.strip() or clean
        parts = [p.strip() for p in TECH_SPLIT_RE.split(clean) if p.strip()]
        if len(parts) == 1 and len(clean.split()) > 10:
            continue
        for part in parts:
            low = part.casefold()
            if low in KNOWN_TECH or len(part) <= 40:
                items.append(part)
    out: list[str] = []
    seen: set[str] = set()
    for item in items:
        key = item.casefold()
        if key not in seen:
            seen.add(key)
            out.append(item)
    return out


def _bucket_for_skill(skill: str) -> str:
    """Map a skill token into a broad presentation bucket."""
    s = skill.casefold()
    if s in {"c#", "java", "python", "typescript", "javascript", "go", "rust", "sql"}:
        return "Languages"
    if s in {"aws", "azure", "gcp"}:
        return "Cloud Platforms"
    if s in {"docker", "kubernetes", "terraform", "github actions", "jenkins", "gitlab ci", "prometheus", "grafana"}:
        return "Tools"
    if s in {"postgresql", "mysql", "redis", "mongodb", "kafka", "rabbitmq"}:
        return "Databases"
    if s in {"microservices", "rest apis", "graphql", "rag", "llm"}:
        return "Other"
    return "Frameworks & Technologies"


def _group_skills(items: list[str]) -> dict[str, list[str]]:
    """Group extracted skills into stable ordered buckets without duplicates."""
    buckets: OrderedDict[str, list[str]] = OrderedDict()
    for item in items:
        bucket = _bucket_for_skill(item)
        buckets.setdefault(bucket, [])
        if item not in buckets[bucket]:
            buckets[bucket].append(item)
    return dict(buckets)


def _normalize_date_line(line: str) -> str:
    """Trim separators and normalize a detected date line."""
    line = _clean_text(line)
    if "·" in line:
        line = line.split("·", 1)[0].strip()
    return line


def _parse_experience(lines: list[str]) -> list[dict[str, Any]]:
    """Parse experience entries using a simple role/company/date line pattern."""
    out: list[dict[str, Any]] = []
    i = 0
    n = len(lines)
    while i < n:
        line = _clean_text(lines[i])
        if not line:
            i += 1
            continue
        if i + 2 < n and DATE_LINE_RE.search(_clean_text(lines[i + 2])):
            role = line
            company = _clean_text(lines[i + 1])
            date_line = _normalize_date_line(lines[i + 2])
            location = ""
            j = i + 3
            # Treat the next short non-bullet line as a probable location.
            if j < n and not BULLET_PREFIX_RE.match(_clean_text(lines[j])) and not DATE_LINE_RE.search(_clean_text(lines[j])):
                candidate = _clean_text(lines[j])
                if len(candidate) <= 80 and not _looks_like_heading(candidate):
                    location = candidate
                    j += 1
            highlights: list[str] = []
            environment: list[str] = []
            while j < n:
                cur = _clean_text(lines[j])
                if not cur:
                    j += 1
                    continue
                if j + 2 < n and DATE_LINE_RE.search(_clean_text(lines[j + 2])) and not BULLET_PREFIX_RE.match(cur):
                    break
                if cur.lower().startswith("environment:") or cur.lower().startswith("tech:"):
                    env_text = cur.split(":", 1)[1].strip() if ":" in cur else cur
                    environment.extend([p.strip() for p in TECH_SPLIT_RE.split(env_text) if p.strip()])
                else:
                    highlights.append(BULLET_PREFIX_RE.sub("", cur))
                j += 1
            out.append({
                "category": "Professional Experience",
                "company_name": company,
                "role": role,
                "dates": {"start": "", "end": "", "display": date_line},
                "location": location,
                "project_description": "",
                "highlights": highlights,
                "environment": list(OrderedDict((x.casefold(), x) for x in environment).values()),
            })
            i = j
            continue
        i += 1
    return out


def _parse_education(lines: list[str]) -> list[dict[str, str]]:
    """Parse education entries as institution / degree / optional year triples."""
    items: list[dict[str, str]] = []
    i = 0
    while i < len(lines):
        institution = _clean_text(lines[i])
        degree = _clean_text(lines[i + 1]) if i + 1 < len(lines) else ""
        year = _clean_text(lines[i + 2]) if i + 2 < len(lines) and DATE_LINE_RE.search(_clean_text(lines[i + 2])) else ""
        if institution:
            items.append({
                "institution": institution,
                "degree": degree,
                "year": year,
                "details": "",
            })
        i += 3 if year else 2
    return items


def docx_baseline_to_qcv_json(baseline: dict[str, Any]) -> dict[str, Any]:
    """Convert deterministic DOCX baseline extraction into approximate QCV JSON."""
    sections = baseline.get("sections") or {}
    preamble = sections.get("preamble") or []
    full_text = baseline.get("full_text") or ""
    contacts = _find_contacts(full_text, baseline.get("links") or [])

    basics = {
        "name": preamble[0] if preamble else "",
        "current_title": preamble[1] if len(preamble) > 1 else "",
        "objective": "",
        "contacts": {
            "email": contacts["email"],
            "phone": contacts["phone"],
            "location": "",
        },
        "links": contacts["links"],
    }

    # If there is no explicit summary section, use a small slice from the preamble.
    summary_lines = sections.get("summary") or []
    if not summary_lines and len(preamble) > 2:
        summary_lines = preamble[2:5]

    skill_lines = []
    for key in ("skills",):
        skill_lines.extend(sections.get(key) or [])

    skills = _group_skills(_tokenize_skills(skill_lines))
    experience = _parse_experience(sections.get("experience") or [])
    education = _parse_education(sections.get("education") or [])
    certifications = [BULLET_PREFIX_RE.sub("", _clean_text(x)) for x in (sections.get("certifications") or []) if _clean_text(x)]
    languages = [{"language": x, "proficiency": "", "level": "", "details": ""} for x in (sections.get("languages") or []) if _clean_text(x)]

    other_sections = []
    if sections.get("skills"):
        other_sections.append({"title": "Top Skills", "items": list(skills.get("Languages", []))[:3]})

    return {
        "basics": basics,
        "summary": {"bullet_points": _split_summary(summary_lines)},
        "skills": skills,
        "experience": experience,
        "education": education,
        "certifications": certifications,
        "languages": languages,
        "other_sections": other_sections,
        "_baseline": {
            "source_type": baseline.get("source_type"),
            "links": baseline.get("links") or [],
        },
    }


def looks_sparse_for_docx(data: dict[str, Any]) -> bool:
    """Heuristic used to detect when deterministic DOCX extraction is too sparse."""
    basics = data.get("basics") or {}
    summary = (data.get("summary") or {}).get("bullet_points") or []
    skills = data.get("skills") or {}
    exp = data.get("experience") or []
    skill_count = sum(len(v) for v in skills.values() if isinstance(v, list))
    return not basics.get("name") or (len(summary) < 1 and skill_count < 3) or len(exp) < 1
